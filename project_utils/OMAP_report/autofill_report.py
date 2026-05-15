import pandas as pd
from docx import Document

# Load your Excel file
excel_file_path = "/Users/jeffreyreina/PycharmProjects/OMAP/OMAP-34-esophagus-CODEX-FF-Casals.xlsx"
df = pd.read_excel(excel_file_path)
df.columns = df.columns.str.strip()  # Remove any whitespace around column names

# Print the column names to troubleshoot
print("Columns in the DataFrame:")
print(df.columns)

# Load your Word document template
template_path = "/Users/jeffreyreina/PycharmProjects/OMAP/avr-template-form_2.docx"

# Loop through each row in the dataframe
for index, row in df.iterrows():
    # Create a new document instance to fill
    filled_doc = Document(template_path)

    # Print values for debugging
    print(f"Row {index} details: {row.to_dict()}")  # Print all row details

    # Define replacements with checks for NaN or empty values
    replacements = {
        "<<TARGET_SYMBOL>>": str(row['target_symbol']),
        "<<HGNC_ID>>": str(row['HGNC_ID']),
        "<<UNIPROT_ACCESSION_NUMBER>>": str(row['uniprot_accession_number']),
        "<<RRID>>": str(row['rrid']),
        "<<HOST>>": str(row['host']),
        "<<ISOTYPE>>": str(row['isotype']),
        "<<CLONALITY>>": str(row['clonality']),
        "<<VENDOR>>": str(row['vendor']),
        "<<CATALOG_NUMBER>>": str(row['catalog_number']),
        "<<RECOMBINANT>>": str(row['recombinant']),
        "<<TISSUE_PRESERVATION>>": str(row['tissue_preservation']),
        "<<ORGAN>>": str(row['organ']),
        "<<ORGAN_UBERON>>": str(row['organ_uberon']),
        "<<METHOD>>": str(row['method']),
        "<<CONJUGATE>>": str(row['conjugate']),
        "<<AUTHOR_ORCIDS>>": str(row['author_orcids']).strip() if pd.notna(row['author_orcids']) else 'NA',
        "<<CELL_TYPES>>": str(row['cell_types']) if pd.notna(row['cell_types']) else 'NA',
        "<<ANATOMICAL_STRUCTURE>>": str(row['anatomical_structure']) if pd.notna(row['anatomical_structure']) else 'NA',
        "<<DS_URL>>": str(row['ds_url']) if pd.notna(row['ds_url']) else 'NA',
        "<<VENDOR_USE>>": str(row['vendor_use']) if pd.notna(row['vendor_use']) else 'NA',
        "<<OMAP_ID>>": str(row['omap_id']),
        "<<LOT_NUMBER>>": str(row['lot_number']),
        "<<DILUTION>>": str(row['dilution_factor']) if pd.notna(row['dilution_factor']) else 'NA',
        "<<ANTIGEN_RETRIEVAL_DETAILS>>": 'NA',  # Fill with 'NA'
        "<<VALIDATION_PROTOCOL_DOI>>": str(row['protocol_doi']),
    }

    # Print replacement values for debugging
    for key, value in replacements.items():
        print(f"Replacing '{key}' with '{value}'")

    # Function to replace text in paragraphs
    def replace_text_in_paragraphs(paragraph, replacements):
        for run in paragraph.runs:
            original_text = run.text
            print(f"Original run text: '{original_text}'")  # Log the original run text
            for placeholder, value in replacements.items():
                if placeholder in original_text:
                    print(f"Replacing '{placeholder}' in: '{original_text}'")  # Debugging line
                    run.text = original_text.replace(placeholder, value)

    # Function to replace text in tables
    def replace_text_in_tables(table, replacements):
        for table_row in table.rows:
            for cell in table_row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraphs(paragraph, replacements)

    # Replace text in paragraphs of the document
    for paragraph in filled_doc.paragraphs:
        replace_text_in_paragraphs(paragraph, replacements)

    # Replace text in text within all tables
    for table in filled_doc.tables:
        replace_text_in_tables(table, replacements)

    # Save the filled document with a unique filename
    filled_doc.save(f"{row['target_symbol']}_OMAP-34.docx")

print("Reports generated successfully!")
